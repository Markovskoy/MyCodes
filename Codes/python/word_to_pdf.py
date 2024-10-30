import pdfplumber
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pytesseract
from pdf2image import convert_from_path

# Путь к PDF и выходному Word файлу
pdf_path = "G:/Диплом/Марковской-В.В.-ОТЧет-НИР-1-семестр.pdf"
word_path = "G:/Диплом/output.docx"

# Инициализация документа Word
document = Document()

# Открываем PDF
with pdfplumber.open(pdf_path) as pdf:
    for page in pdf.pages:
        # Извлечение текста страницы с обработкой таблиц
        text = page.extract_text()
        
        if text:
            # Разделение текста на абзацы и добавление в документ
            for paragraph in text.split('\n'):
                if paragraph.strip():
                    p = document.add_paragraph(paragraph.strip())
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    run = p.runs[0]
                    run.font.size = Pt(12)  # Примерный размер шрифта

        # Извлечение таблиц и вставка
    tables = page.extract_tables()
    for table_data in tables:
        table = document.add_table(rows=1, cols=len(table_data[0]))
        table.style = 'Table Grid'
        for row in table_data:
            cells = table.add_row().cells
            for idx, cell_text in enumerate(row):
                # Проверяем, что cell_text не None
                cells[idx].text = cell_text if cell_text is not None else ""

# OCR обработка изображений страниц, если PDF содержит не текст, а сканы
images = convert_from_path(pdf_path)
for img in images:
    text = pytesseract.image_to_string(img, lang='rus+eng')
    if text.strip():
        document.add_paragraph(text)

# Сохранение документа
document.save(word_path)
print(f"Word document saved at {word_path}")